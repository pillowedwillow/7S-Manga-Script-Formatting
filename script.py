import docx
from docx.shared import Pt
# tkinter handles the gui
import tkinter as tk
from tkinter import simpledialog


ROOT = tk.Tk()
ROOT.withdraw()
# the input dialog
user_inp = simpledialog.askstring(title="7S Manga Template", prompt="How many pages?")
 
# open document
mydoc = docx.Document("Core.docx")

# convert input string into an integer for range()
a = int(user_inp)

for y in range(a):
    page = mydoc.add_paragraph("")
    page_format = page.paragraph_format
    page_format.space_after = Pt(0)
    page_format.line_spacing = 1
    # page number
    words = ["Page ", str(y+1)]
    run = page.add_run(words)
    run.bold = True
    run.underline = True


    collection = [0.1, 0.2, 0.3, 0.4, 0.5, 0.6, 0.7, 0.8]
    for x in collection:
        panel = mydoc.add_paragraph("")
        panel_format = panel.paragraph_format
        panel_format.space_after = Pt(0)
        panel_format.line_spacing = 1
        # panel number
        c = y + 1 + x
        run = panel.add_run(str(c))
        run.bold = True
        # space after
        empty = mydoc.add_paragraph("")
        empty_format = empty.paragraph_format
        empty_format.space_after = Pt(0)
        empty_format.line_spacing = 1

    # space after
    empty = mydoc.add_paragraph("")
    empty_format = empty.paragraph_format
    empty_format.space_after = Pt(0)
    empty_format.line_spacing = 1

# save document
mydoc.save("7S Manga Script.docx")
